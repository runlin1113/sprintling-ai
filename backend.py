import cv2
import numpy as np
import json
import matplotlib.pyplot as plt
from scipy.signal import find_peaks
from ultralytics import YOLO
import torch
import os
from docx import Document
import ai_coach

_model = None


def get_yolo_model():
    global _model
    if _model is None:
        _model = YOLO('yolov8n-pose.pt')
        _model.to('cuda' if torch.cuda.is_available() else 'cpu')
    return _model


def calculate_angle(a, b, c):
    a, b, c = np.array(a), np.array(b), np.array(c)
    radians = np.arctan2(c[1] - b[1], c[0] - b[0]) - np.arctan2(a[1] - b[1], a[0] - b[0])
    angle = np.abs(radians * 180.0 / np.pi)
    return 360 - angle if angle > 180.0 else angle


def calculate_torso_lean(shoulder, hip):
    dx = shoulder[0] - hip[0]
    dy = hip[1] - shoulder[1]
    angle_rad = np.arctan2(dx, dy)
    return round(angle_rad * 180.0 / np.pi, 2)


def draw_bone(img, pt1, pt2, color, thickness):
    if pt1 != (0, 0) and pt2 != (0, 0):
        cv2.line(img, pt1, pt2, color, thickness, cv2.LINE_AA)


def process_full_kinematics(video_path, output_video, output_json):
    """独立处理单段视频并提取运动学特征"""
    model = get_yolo_model()
    cap = cv2.VideoCapture(video_path)

    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    if fps == 0 or np.isnan(fps): fps = 30.0

    fourcc = cv2.VideoWriter_fourcc(*'vp09')
    out = cv2.VideoWriter(output_video, fourcc, fps, (width, height))

    timeseries_data = []
    frame_count = 0
    prev_time, prev_com_x, prev_r_knee_angle = 0.0, 0, 0.0
    prev_kpts = {i: (0, 0) for i in range(17)}
    EMA_ALPHA = 0.85
    COLOR_RED = (0, 0, 255)

    while cap.isOpened():
        ret, frame = cap.read()
        if not ret: break

        frame_count += 1
        results = model.predict(source=frame, verbose=False, conf=0.25, imgsz=1024)

        frame_metrics = {
            "time_sec": round(frame_count / fps, 3),
            "torso_lean": 0.0,
            "angles": {},
            "kinetics": {"com_velocity_x": 0.0, "r_knee_angular_velocity": 0.0, "head_posture_deviation": 0.0}
        }

        if results[0].keypoints is not None and results[0].boxes is not None and len(results[0].keypoints) > 0:
            boxes = results[0].boxes.xywh.cpu().numpy()
            areas = boxes[:, 2] * boxes[:, 3]
            if len(areas) > 0:
                target_idx = np.argmax(areas)
                raw_kpts = results[0].keypoints.data[target_idx].cpu().numpy()

                try:
                    smoothed_kpts = {}
                    for idx in range(17):
                        x, y = int(raw_kpts[idx][0]), int(raw_kpts[idx][1])
                        if x != 0 and y != 0:
                            if prev_kpts[idx] != (0, 0):
                                smooth_x = int(EMA_ALPHA * x + (1 - EMA_ALPHA) * prev_kpts[idx][0])
                                smooth_y = int(EMA_ALPHA * y + (1 - EMA_ALPHA) * prev_kpts[idx][1])
                                smoothed_kpts[idx] = (smooth_x, smooth_y)
                            else:
                                smoothed_kpts[idx] = (x, y)
                            prev_kpts[idx] = smoothed_kpts[idx]
                        else:
                            smoothed_kpts[idx] = (0, 0)
                            prev_kpts[idx] = (0, 0)

                    head_pt = smoothed_kpts[0]
                    l_sh, r_sh = smoothed_kpts[5], smoothed_kpts[6]
                    l_el, r_el = smoothed_kpts[7], smoothed_kpts[8]
                    l_hip, r_hip = smoothed_kpts[11], smoothed_kpts[12]
                    l_kn, r_kn = smoothed_kpts[13], smoothed_kpts[14]
                    l_an, r_an = smoothed_kpts[15], smoothed_kpts[16]

                    if l_sh[0] and r_hip[0]: frame_metrics["torso_lean"] = calculate_torso_lean(r_sh, r_hip)

                    com_x = int((l_hip[0] + r_hip[0]) / 2) if l_hip[0] and r_hip[0] else 0

                    if all(pt[0] for pt in [r_hip, r_kn, r_an]): frame_metrics["angles"]["r_knee"] = round(
                        calculate_angle(r_hip, r_kn, r_an), 2)
                    if all(pt[0] for pt in [l_hip, l_kn, l_an]): frame_metrics["angles"]["l_knee"] = round(
                        calculate_angle(l_hip, l_kn, l_an), 2)

                    if head_pt[0] and r_sh[0] and r_hip[0]:
                        triangle_area = abs((head_pt[0] * (r_sh[1] - r_hip[1]) + r_sh[0] * (r_hip[1] - head_pt[1]) +
                                             r_hip[0] * (head_pt[1] - r_sh[1])) / 2.0)
                        frame_metrics["kinetics"]["head_posture_deviation"] = round(triangle_area, 2)

                    current_time = frame_metrics["time_sec"]
                    dt = current_time - prev_time
                    if dt > 0 and prev_time > 0 and com_x > 0:
                        dx = com_x - prev_com_x
                        frame_metrics["kinetics"]["com_velocity_x"] = round(dx / dt, 2)
                        current_angle = frame_metrics["angles"].get("r_knee", 0)
                        if current_angle and prev_r_knee_angle:
                            frame_metrics["kinetics"]["r_knee_angular_velocity"] = round(
                                (current_angle - prev_r_knee_angle) / dt, 2)

                    prev_time = current_time
                    prev_com_x = com_x
                    prev_r_knee_angle = frame_metrics["angles"].get("r_knee", 0)

                    bones = [(l_sh, l_el), (l_el, smoothed_kpts[9]), (l_hip, l_kn), (l_kn, l_an),
                             (r_sh, r_hip), (r_sh, r_el), (r_el, smoothed_kpts[10]), (r_hip, r_kn), (r_kn, r_an),
                             (l_sh, r_sh), (l_hip, r_hip)]
                    for pt1, pt2 in bones: draw_bone(frame, pt1, pt2, COLOR_RED, 3)
                    for pt in [l_sh, l_el, smoothed_kpts[9], l_hip, l_kn, l_an, r_sh, r_el, smoothed_kpts[10], r_hip,
                               r_kn, r_an]:
                        if pt[0] != 0: cv2.circle(frame, pt, 4, COLOR_RED, -1, cv2.LINE_AA)
                except Exception:
                    pass

        timeseries_data.append(frame_metrics)
        out.write(frame)

    cap.release()
    out.release()
    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(timeseries_data, f, ensure_ascii=False, indent=4)

    if torch.cuda.is_available(): torch.cuda.empty_cache()


def extract_combined_features(start_json, maxvel_json, output_payload):
    """联合处理数据，动态判断存在的模态并安全提取"""

    def safe_mean(arr):
        return round(np.mean(arr), 2) if len(arr) > 0 else 0

    def safe_max(arr):
        return round(np.max(arr), 2) if len(arr) > 0 else 0

    payload = {"metadata": {"task": "Sprint Biomechanics Analysis"}}

    # 动态解析起跑期数据 (Phase 1)
    if os.path.exists(start_json):
        with open(start_json, 'r') as f: start_data = json.load(f)
        com_vel_start = [d['kinetics']['com_velocity_x'] for d in start_data if 'kinetics' in d]
        head_dev_start = [d['kinetics']['head_posture_deviation'] for d in start_data if 'kinetics' in d]
        payload["phase_1_start_acceleration"] = {
            "max_com_velocity_px_s": safe_max(com_vel_start),
            "average_head_deviation": safe_mean(head_dev_start),
            "torso_lean_initial": start_data[0]['torso_lean'] if start_data else 0
        }

    # 动态解析极速期数据 (Phase 2)
    if os.path.exists(maxvel_json):
        with open(maxvel_json, 'r') as f: maxvel_data = json.load(f)
        l_knees = [d['angles']['l_knee'] for d in maxvel_data if d['angles'].get('l_knee', 0) > 0]
        r_knees = [d['angles']['r_knee'] for d in maxvel_data if d['angles'].get('r_knee', 0) > 0]
        r_knee_ang_vel = [abs(d['kinetics']['r_knee_angular_velocity']) for d in maxvel_data if 'kinetics' in d]

        r_peaks, _ = find_peaks(-np.array(r_knees), distance=10, prominence=10)
        l_peaks, _ = find_peaks(-np.array(l_knees), distance=10, prominence=10)

        r_max_flex = safe_mean(np.array(r_knees)[r_peaks]) if len(r_peaks) > 0 else 0
        l_max_flex = safe_mean(np.array(l_knees)[l_peaks]) if len(l_peaks) > 0 else 0
        si = abs(l_max_flex - r_max_flex) / (0.5 * (l_max_flex + r_max_flex)) * 100 if (
                                                                                                   l_max_flex + r_max_flex) > 0 else 0

        r_cv = round(np.std(r_knees) / np.mean(r_knees), 4) if len(r_knees) > 0 and np.mean(r_knees) > 0 else 0

        payload["phase_2_maximum_velocity"] = {
            "max_knee_angular_velocity": safe_max(r_knee_ang_vel),
            "flexion_symmetry_index": round(si, 2),
            "cns_stability_coefficient": r_cv
        }

    with open(output_payload, 'w', encoding='utf-8') as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def plot_combined_dashboard(start_json, maxvel_json, output_img):
    """根据实际存在的数据，动态绘制1个或2个图表"""
    has_start = os.path.exists(start_json)
    has_max = os.path.exists(maxvel_json)

    if not has_start and not has_max:
        return

    num_plots = sum([has_start, has_max])
    plt.style.use('dark_background')
    fig, axes = plt.subplots(num_plots, 1, figsize=(10, 3.5 * num_plots), dpi=120)
    fig.patch.set_facecolor('#1A1A1A')

    # 兼容单图表返回对象不是列表的情况
    if num_plots == 1:
        axes = [axes]

    plot_idx = 0

    if has_start:
        ax = axes[plot_idx]
        with open(start_json, 'r') as f: start_data = json.load(f)
        times_s = [d['time_sec'] for d in start_data]
        leans = [d['torso_lean'] if d['torso_lean'] > 0 else np.nan for d in start_data]
        ax.plot(times_s, leans, color='#D32F2F', linewidth=2)
        ax.set_title('Phase 1 (Start): Torso Lean Dynamics', color='#E0E0E0', fontsize=12)
        ax.set_ylabel('Angle (deg)')
        ax.grid(True, linestyle=':', alpha=0.3)
        plot_idx += 1

    if has_max:
        ax = axes[plot_idx]
        with open(maxvel_json, 'r') as f: maxvel_data = json.load(f)
        times_m = [d['time_sec'] for d in maxvel_data]
        l_k = [d['angles'].get('l_knee', np.nan) for d in maxvel_data]
        r_k = [d['angles'].get('r_knee', np.nan) for d in maxvel_data]
        ax.plot(times_m, l_k, color='#9E9E9E', label='Left Knee')
        ax.plot(times_m, r_k, color='#D32F2F', label='Right Knee')
        ax.set_title('Phase 2 (Max Vel): Knee Flexion Symmetry', color='#E0E0E0', fontsize=12)
        ax.set_ylabel('Flexion (deg)')
        ax.grid(True, linestyle=':', alpha=0.3)
        ax.legend(loc="upper right")

    plt.tight_layout()
    plt.savefig(output_img, facecolor=fig.get_facecolor(), edgecolor='none')
    plt.close(fig)


def create_docx_report(md_filepath, docx_filepath):
    if not os.path.exists(md_filepath): return
    doc = Document()
    doc.add_heading('Biomechanical Diagnosis Report', 0)

    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line: continue
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(docx_filepath)


def generate_training_report(payload_path, output_md, athlete_info=None):
    return ai_coach.generate_training_prescription(payload_path, output_md, athlete_info)